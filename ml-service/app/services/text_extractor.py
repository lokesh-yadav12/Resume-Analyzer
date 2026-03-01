import fitz  # PyMuPDF
from docx import Document
import logging
import os

logger = logging.getLogger(__name__)

class TextExtractor:
    """Service for extracting text from various file formats"""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._extract_pdf_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'text/plain': self._extract_txt_text
        }
    
    def extract_text(self, file_path, mime_type):
        """
        Extract text from file based on mime type
        
        Args:
            file_path (str): Path to the file
            mime_type (str): MIME type of the file
            
        Returns:
            str: Extracted text content
            
        Raises:
            ValueError: If file type is not supported
            Exception: If extraction fails
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if mime_type not in self.supported_types:
            raise ValueError(f"Unsupported file type: {mime_type}")
        
        try:
            extractor_func = self.supported_types[mime_type]
            text = extractor_func(file_path)
            
            if not text or len(text.strip()) < 10:
                raise Exception("No readable text found in the file")
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text: {str(e)}")
    
    def _extract_pdf_text(self, file_path):
        """Extract text from PDF file using PyMuPDF"""
        try:
            text = ""
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                
                # Clean up the text
                page_text = self._clean_text(page_text)
                text += page_text + "\n"
            
            doc.close()
            
            # If no text extracted, try OCR-like extraction
            if len(text.strip()) < 50:
                text = self._extract_pdf_with_ocr_fallback(file_path)
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise Exception(f"Failed to extract PDF text: {str(e)}")
    
    def _extract_pdf_with_ocr_fallback(self, file_path):
        """Fallback method for PDFs that might be image-based"""
        try:
            text = ""
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Try to get text blocks with more detailed extraction
                blocks = page.get_text("dict")
                
                for block in blocks.get("blocks", []):
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line.get("spans", []):
                                text += span.get("text", "") + " "
                            text += "\n"
            
            doc.close()
            return text
            
        except Exception as e:
            logger.error(f"OCR fallback failed: {str(e)}")
            return ""
    
    def _extract_docx_text(self, file_path):
        """Extract text from DOCX file using python-docx"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            # Extract text from headers and footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise Exception(f"Failed to extract DOCX text: {str(e)}")
    
    def _extract_txt_text(self, file_path):
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        return self._clean_text(text)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
                return self._clean_text(text)
                
        except Exception as e:
            logger.error(f"Error extracting TXT text: {str(e)}")
            raise Exception(f"Failed to extract TXT text: {str(e)}")
    
    def _clean_text(self, text):
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        import re
        
        # Replace multiple spaces with single space
        text = re.sub(r'\s+', ' ', text)
        
        # Replace multiple newlines with double newline
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove excessive punctuation
        text = re.sub(r'[.]{3,}', '...', text)
        text = re.sub(r'[-]{3,}', '---', text)
        
        # Fix common OCR errors
        text = text.replace('|', 'I')  # Common OCR mistake
        text = text.replace('0', 'O')  # In names/words, not numbers
        
        # Remove control characters
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        
        return text.strip()
    
    def get_file_info(self, file_path, mime_type):
        """Get basic information about the file"""
        try:
            file_size = os.path.getsize(file_path)
            
            info = {
                'file_size': file_size,
                'mime_type': mime_type,
                'pages': 0,
                'has_images': False,
                'has_tables': False
            }
            
            if mime_type == 'application/pdf':
                doc = fitz.open(file_path)
                info['pages'] = len(doc)
                
                # Check for images and tables (simplified)
                for page_num in range(min(3, len(doc))):  # Check first 3 pages
                    page = doc.load_page(page_num)
                    images = page.get_images()
                    if images:
                        info['has_images'] = True
                    
                    # Simple table detection based on text structure
                    text = page.get_text()
                    if '\t' in text or '|' in text:
                        info['has_tables'] = True
                
                doc.close()
            
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                doc = Document(file_path)
                info['pages'] = len(doc.paragraphs) // 20  # Rough estimate
                
                if doc.tables:
                    info['has_tables'] = True
                
                # Check for images (simplified)
                for paragraph in doc.paragraphs[:10]:  # Check first 10 paragraphs
                    if paragraph.runs:
                        for run in paragraph.runs:
                            if hasattr(run, '_element') and run._element.xpath('.//pic:pic'):
                                info['has_images'] = True
                                break
            
            return info
            
        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            return {
                'file_size': 0,
                'mime_type': mime_type,
                'pages': 0,
                'has_images': False,
                'has_tables': False
            }